"""
Controller for AI-based exam generation.

This module orchestrates:

1. Input validation.
2. Class and exercise type validation.
3. Exam creation in GENERATING state.
4. RAG retrieval from Qdrant.
5. LLM prompt construction and generation.
6. JSON validation of LLM output.
7. Persistence of generated snapshot.
"""

import json
import datetime
from typing import cast, List, Any, Dict
import io
from flask import request, jsonify, send_file
from flask_jwt_extended import get_jwt_identity, jwt_required
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import joinedload
from sqlalchemy import or_
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.platypus import ListFlowable, ListItem
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import KeepTogether
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from orm_models import db, Exam, Class, Exercise, ExamExerciseInstance, User
from services.rag_service import retrieve_course_context
from services.llm_service import build_prompt, generate_exam_from_llm, build_refinement_prompt, build_student_prompt, build_student_correction_prompt, generate_student_correction_from_llm
from services.generic_context_service import get_random_generic_context
from utils.email_utils import send_exam_accepted_email_to_teacher, send_exam_corrected_email_to_coordinator, send_exam_on_review_email_to_teacher, send_exam_sent_to_correction_email_to_teacher
from utils.types_enum import ExamStatus
from utils.exam_xp import calculate_exam_xp, resolve_level_from_xp, apply_exam_xp_to_student
from utils.mpreg_utils import predict_next_student_score, get_difficulty_band
from services.exam_fallback_service import ExamFallbackService, StudentExamFallbackService
import time
import math

def extract_json(text: str) -> str:
    """
    Docstring for extract_json
    
    :param text: Description
    :type text: str
    :return: Description
    :rtype: str
    """
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("No JSON object found in model output")
    return text[start:end + 1]

@jwt_required()
def generate_exam():
    """
    Generate a new exam using RAG + LLM.

    Expected JSON payload:
    {
        "class_id": int,
        "context": str,
        "exercise_type_ids": [int]
    }

    Returns:
        201 with exam_id if successful.
    """

    data = request.get_json(silent=True)
    if not data:
        return jsonify({"message": "Invalid JSON body"}), 400

    class_id = data.get("class_id")
    context = data.get("context")
    exercise_type_ids = data.get("exercise_type_ids")
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)

    if not class_id or not context or not exercise_type_ids:
        return jsonify({"message": "Missing required fields"}), 400

    if not isinstance(exercise_type_ids, list) or not exercise_type_ids:
        return jsonify({"message": "exercise_type_ids must be a non-empty list"}), 400
    if not user:
        return jsonify({"message":"Teacher not found"}), 404

    try:
        # ------------------------
        # Validate Class
        # ------------------------
        class_obj = Class.query.get(class_id)
        if not class_obj or class_obj.date_deleted:
            return jsonify({"message": "Class not found or deleted"}), 404

        # ------------------------
        # Validate Exercise Types
        # ------------------------
        exercise_types: List[Exercise] = []

        for ex_id in exercise_type_ids:
            ex = Exercise.query.get(ex_id)
            if not ex or ex.date_deleted:
                return jsonify(
                    {"message": f"Invalid exercise type id {ex_id}"}
                ), 400
            exercise_types.append(ex)

        # ------------------------
        # Create Exam (GENERATING)
        # ------------------------
        new_exam = Exam(
            status=ExamStatus.GENERATING.value,
            class_id=class_id,
            context=context,
            user_id=user_id
        )

        for ex in exercise_types:
            new_exam.exercise_types.append(ex)

        db.session.add(new_exam)
        db.session.flush()  # ensures new_exam.id exists

        # ------------------------
        # Generation Flow / Fallback
        # ------------------------
        level = class_obj.suggested_level

        try:
            # ------------------------
            # RAG Phase
            # ------------------------
            exercise_list_text = "\n".join(
                [f"- {ex.name}: {ex.content_description}" for ex in exercise_types]
            )

            contexts = retrieve_course_context(
                course_id=class_obj.description,
                level=level,
                exercises_description=exercise_list_text,
            )

            retrieved_context_text = "\n\n---\n\n".join(contexts)

            # ------------------------
            # LLM Phase
            # ------------------------
            prompt = build_prompt(
                level=level,
                teacher_text=context,
                exercise_list_text=exercise_list_text,
                retrieved_context=retrieved_context_text,
            )

            raw_output = generate_exam_from_llm(prompt)

            #print("----- MODEL RAW OUTPUT -----")
            #print(raw_output)
            #print("----- END MODEL OUTPUT -----")

            cleaned_json = extract_json(raw_output)
            parsed_output = json.loads(cleaned_json)

            if "exercises" not in parsed_output:
                raise ValueError("Invalid exam structure returned by model")

        except Exception as err:
            print("Generation flow failed, using fallback:", err)
            time.sleep(10)
            parsed_output = ExamFallbackService.build_exam_payload(
                level=level,
                exercise_types=exercise_types,
            )

            raw_output = ExamFallbackService.build_snapshot(
                reason=str(err),
                payload=parsed_output,
            )

        # ------------------------
        # Persist Exercise Instances
        # ------------------------
        for exercise_block in parsed_output["exercises"]:
            exercise_name = exercise_block["exercise_type"]

            # buscar el exercise type en BD
            exercise_type = Exercise.query.filter(
                db.func.lower(Exercise.name) == exercise_name.lower()
            ).first()

            if not exercise_type:
                db.session.rollback()
                return jsonify({
                    "message": f"Exercise type '{exercise_name}' not found in catalog"
                }), 400

            instance = ExamExerciseInstance(
                exam_id=new_exam.id,
                exercise_type_id=exercise_type.id,
                instructions=exercise_block["instructions"],
                content_json=json.dumps(exercise_block["items"]),
                answer_key_json=json.dumps({
                    "answers": ExamFallbackService.extract_answers(
                        exercise_block["items"]
                    )
                })
            )
            db.session.add(instance)

        # ------------------------
        # Persist Results
        # ------------------------
        new_exam.generated_snapshot = raw_output
        new_exam.status = ExamStatus.GENERATING.value

        db.session.commit()

        return jsonify(
            {
                "message": "Exam generated successfully",
                "exam_id": new_exam.id,
            }
        ), 201
    except SQLAlchemyError as err:
        print('Database error', err)
        db.session.rollback()
        return jsonify({"message": f"Database error: {err}"}), 500

    except Exception as err:  # pylint: disable=broad-except
        print('Unexpected err', err)
        db.session.rollback()
        return jsonify({"message": f"Unexpected error: {err}"}), 500

def refine_exam(exam_id: int):
    """
    Docstring for refine_exam
    
    :param exam_id: Description
    :type exam_id: int
    """
    data = request.get_json(silent=True)
    if not data:
        return jsonify({"message": "Invalid JSON body"}), 400

    feedback = data.get("feedback")
    if not feedback:
        return jsonify({"message": "Missing feedback"}), 400

    exam = Exam.query.get(exam_id)
    if not exam or exam.date_deleted:
        return jsonify({"message": "Exam not found"}), 404

    original_status = exam.status
    if int(exam.user_id) != int(get_jwt_identity()):
        return jsonify({"message": "Unauthorized"}), 403

    if original_status not in (ExamStatus.GENERATING.value, ExamStatus.ON_CORRECTION.value):
        return jsonify({"message": "Exam is not editable in current state"}), 400

    try:
        # Nivel para el prompt (mismo criterio que en generate)
        level = exam.class_exam.suggested_level

        refinement_prompt = build_refinement_prompt(
            level=level,
            original_snapshot=exam.generated_snapshot,
            teacher_feedback=feedback,
        )
        raw_output = generate_exam_from_llm(refinement_prompt)

        cleaned_json = extract_json(raw_output)
        parsed_output = json.loads(cleaned_json)

        if "exercises" not in parsed_output:
            return jsonify({"message": "Invalid exam structure returned by model"}), 500

        # Validar que no cambien los exercise types originales
        original_types = {ex.name.lower() for ex in exam.exercise_types}
        returned_types = {
            ex_block["exercise_type"].lower()
            for ex_block in parsed_output["exercises"]
        }

        if original_types != returned_types:
            return jsonify({
                "message": "Refinement cannot change exercise types"
            }), 400

        # borrar instancias previas
        ExamExerciseInstance.query.filter_by(exam_id=exam.id).delete()

        for exercise_block in parsed_output["exercises"]:
            exercise_type = Exercise.query.filter(
                db.func.lower(Exercise.name)
                == exercise_block["exercise_type"].lower()
            ).first()

            # FIX Pylance: validar None explícitamente
            if not exercise_type:
                db.session.rollback()
                return jsonify({
                    "message": f"Exercise type '{exercise_block['exercise_type']}' not found"
                }), 400

            instance = ExamExerciseInstance(
                exam_id=exam.id,
                exercise_type_id=exercise_type.id,
                instructions=exercise_block["instructions"],
                content_json=json.dumps(exercise_block["items"]),
                answer_key_json=json.dumps({
                    "answers": [
                        item["answer"] for item in exercise_block["items"]
                    ]
                })
            )
            db.session.add(instance)

        exam.generated_snapshot = raw_output
        if original_status == ExamStatus.GENERATING.value:
            exam.status = ExamStatus.GENERATING.value
        else:
            exam.status = ExamStatus.ON_CORRECTION.value

        db.session.commit()

        return jsonify({"message": "Exam refined successfully"}), 200

    except Exception as err:  # pylint: disable=broad-except
        db.session.rollback()
        return jsonify({"message": str(err)}), 500

def get_full_exam(exam_id: int):
    """
    Return a fully reconstructed exam ready for frontend rendering.
    """

    try:
        exam = (
            Exam.query
            .options(
                joinedload(Exam.class_exam), # type: ignore
                joinedload(Exam.user_exam),          # teacher # type: ignore
                joinedload(Exam.coordinator_exam),  # coordinator # type: ignore
                joinedload(Exam.generated_exercises) # type: ignore
            )
            .filter(
                Exam.id == exam_id,
                Exam.date_deleted.is_(None)
            )
            .first()
        )

        if not exam:
            return jsonify({"message": "Exam not found"}), 404

        result = {
            "id": exam.id,
            "status": exam.status,
            "context": exam.context,
            "notes": exam.notes,
            "date_created": exam.date_created,

            # Clase
            "class_id": exam.class_id,
            "class_description": (
                exam.class_exam.description
                if exam.class_exam else None
            ),

            # Profesora
            "teacher_id": exam.user_id,
            "teacher_full_name": (
                f"{exam.user_exam.name} {exam.user_exam.surname}"
                if exam.user_exam else None
            ),

            # Coordinadora
            "coordinator_id":exam.coordinator_id,
            "coordinator_full_name": (
                f"{exam.coordinator_exam.name} {exam.coordinator_exam.surname}"
                if exam.coordinator_exam else None
            ),

            "exercises": []
        }

        for instance in exam.generated_exercises:
            exercise_type = instance.exercise_type

            result["exercises"].append({
                "exam_exercise_instance_id": instance.id,
                "exercise_type": exercise_type.name if exercise_type else None,
                "instructions": instance.instructions,
                "items": json.loads(instance.content_json)
            })

        return jsonify(result), 200

    except Exception as err:
        return jsonify({"message": f"Unexpected error: {err}"}), 500

def build_exam_html(exam: Exam) -> str:
    """
    Build HTML representation of exam for PDF export.
    """

    html = """
    <html>
    <head>
        <style>
            body { font-family: Arial, sans-serif; }
            h1 { text-align: center; }
            h2 { margin-top: 30px; }
            .question { margin-bottom: 10px; }
            .options { margin-left: 20px; }
            .page-break { page-break-before: always; }
        </style>
    </head>
    <body>
    """

    html += "<h1>Exam</h1>"
    html += f"<p><strong>Class ID:</strong> {exam.class_id}</p>"
    html += f"<p><strong>Date:</strong> {exam.date_created}</p>"
    html += "<hr>"

    generated_exercises = cast(List[ExamExerciseInstance], exam.generated_exercises)
    for instance in generated_exercises:
        items = json.loads(instance.content_json)

        html += f"<h2>{instance.exercise_type.name}</h2>"
        html += f"<p><em>{instance.instructions}</em></p>"

        for idx, item in enumerate(items, start=1):
            html += f"<div class='question'><strong>{idx}.</strong> {item['question']}</div>"

            if "options" in item:
                for option in item["options"]:
                    html += f"<div class='options'>{option}</div>"

        html += "<br>"

    # Answer sheet page
    html += "<div class='page-break'></div>"
    html += "<h1>Answer Sheet</h1>"

    generated_exercises = cast(List[ExamExerciseInstance], exam.generated_exercises)
    for instance in generated_exercises:
        answers = json.loads(instance.answer_key_json)["answers"]
        html += f"<h3>{instance.exercise_type.name}</h3>"
        for idx, ans in enumerate(answers, start=1):
            html += f"<p>{idx}. {ans}</p>"

    html += "</body></html>"

    return html


def export_exam_pdf(exam_id: int):
    """
    Generate a PDF version of a generated exam using ReportLab.
    """

    exam = Exam.query.get(exam_id)
    if not exam or exam.date_deleted:
        return jsonify({"message": "Exam not found"}), 404

    if exam.status != ExamStatus.ACCEPTED.value:
        return jsonify({"message": "Exam not generated yet"}), 400

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    elements = []

    styles = getSampleStyleSheet()
    title_style = styles["Heading1"]
    normal_style = styles["Normal"]

    elements.append(Paragraph("Exam", title_style))
    elements.append(Spacer(1, 0.3 * inch))

    for instance in exam.generated_exercises:
        elements.append(Paragraph(instance.exercise_type.name, styles["Heading2"]))
        elements.append(Spacer(1, 0.2 * inch))

        elements.append(Paragraph(instance.instructions, normal_style))
        elements.append(Spacer(1, 0.2 * inch))

        items = json.loads(instance.content_json)

        for idx, item in enumerate(items, start=1):
            question_text = f"{idx}. {item['question']}"
            elements.append(Paragraph(question_text, normal_style))
            elements.append(Spacer(1, 0.1 * inch))

            if "options" in item:
                for option in item["options"]:
                    elements.append(Paragraph(option, normal_style))
                    elements.append(Spacer(1, 0.1 * inch))

        elements.append(Spacer(1, 0.3 * inch))

    # Answer sheet
    elements.append(PageBreak())
    elements.append(Paragraph("Answer Sheet", title_style))
    elements.append(Spacer(1, 0.3 * inch))

    for instance in exam.generated_exercises:
        elements.append(Paragraph(instance.exercise_type.name, styles["Heading2"]))
        elements.append(Spacer(1, 0.2 * inch))

        answers = json.loads(instance.answer_key_json)["answers"]

        for idx, answer in enumerate(answers, start=1):
            elements.append(Paragraph(f"{idx}. {answer}", normal_style))
            elements.append(Spacer(1, 0.1 * inch))

        elements.append(Spacer(1, 0.3 * inch))

    doc.build(elements)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"exam_{exam_id}.pdf",
        mimetype="application/pdf"
    )

def export_exam_docx(exam_id: int):
    """
    Generate a DOCX version of a generated exam.
    """

    exam = Exam.query.get(exam_id)
    if not exam or exam.date_deleted:
        return jsonify({"message": "Exam not found"}), 404

    if exam.status != ExamStatus.ACCEPTED.value:
        return jsonify({"message": "Exam not generated yet"}), 400

    document = Document()

    # Title
    document.add_heading("Exam", level=1)

    document.add_paragraph(f"Class ID: {exam.class_id}")
    document.add_paragraph(f"Date: {exam.date_created}")
    document.add_paragraph("\n")

    # Exercises
    for instance in exam.generated_exercises:
        document.add_heading(instance.exercise_type.name, level=2)
        document.add_paragraph(instance.instructions)

        items = json.loads(instance.content_json)

        for idx, item in enumerate(items, start=1):
            document.add_paragraph(f"{idx}. {item['question']}")

            if "options" in item:
                for option in item["options"]:
                    document.add_paragraph(option, style="List Bullet")

        document.add_paragraph("\n")

    # Answer Sheet
    document.add_page_break()
    document.add_heading("Answer Sheet", level=1)

    for instance in exam.generated_exercises:
        document.add_heading(instance.exercise_type.name, level=2)

        answers = json.loads(instance.answer_key_json)["answers"]

        for idx, answer in enumerate(answers, start=1):
            document.add_paragraph(f"{idx}. {answer}")

        document.add_paragraph("\n")

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"exam_{exam_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def delete_exam(exam_id: int):
    """
    Docstring for delete_exam
    
    :param exam_id: Description
    :type exam_id: int
    """
    exam = Exam.query.get(exam_id)
    if not exam or exam.date_deleted:
        return jsonify({"message": "Exam not found"}), 404
    try:
        exam.date_deleted = datetime.datetime.now()
        db.session.commit()
        return jsonify({"message":f"Level {exam.id} deleted successfully"}), 200
    except SQLAlchemyError as err:
        db.session.rollback()
        return jsonify({"message": f"Database error: {err}"}), 500
    except Exception as err:  # pylint: disable=broad-except
        db.session.rollback()
        return jsonify({"message": {f"Something went wrong: {err}"}}), 500

def get_all_exams_controller():
    """
    Docstring for get_all_exams_controller
    """
    try:
        current_user_id = get_jwt_identity()

        exams = (
            Exam.query
            .options(
                joinedload(Exam.class_exam),   # type: ignore
                joinedload(Exam.user_exam)     # type: ignore
            )
            .filter(
                Exam.date_deleted.is_(None),
                Exam.status != ExamStatus.GENERATING.value,
                Exam.status != ExamStatus.STUDENT_EXAM.value,
                Exam.status != ExamStatus.SOLVED.value,
                Exam.status != ExamStatus.TEST_EXAM.value,
                or_(
                    Exam.coordinator_id.is_(None),
                    Exam.coordinator_id == current_user_id
                )
            )
            .order_by(Exam.date_created.desc())
            .all()
        )

        result = []
        for exam in exams:
            result.append({
                "id": exam.id,
                "status": exam.status,
                "context": exam.context,
                "class_id": exam.class_id,
                "class_description": exam.class_exam.description if exam.class_exam else None,
                "user_id": exam.user_id,
                "teacher_full_name": (
                    f"{exam.user_exam.name} {exam.user_exam.surname}"
                    if exam.user_exam else None
                ),
                "generated_exercises": [],
                "date_created": exam.date_created,
                "coordinator_id":exam.coordinator_id
            })

        return jsonify(result), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@jwt_required()
def get_teacher_exams_controller():
    """
    Docstring for get_teacher_exams_controller
    """
    try:
        current_user_id = int(get_jwt_identity())
        exams = (
            Exam.query
            .options(
                joinedload(Exam.class_exam),         # type: ignore
                joinedload(Exam.coordinator_exam)    # type: ignore
            )
            .filter(
                Exam.date_deleted.is_(None),
                Exam.user_id == current_user_id,
                Exam.status != ExamStatus.GENERATING.value
            )
            .order_by(Exam.date_created.desc())
            .all()
        )

        result = []
        for exam in exams:
            result.append({
                "id": exam.id,
                "status": exam.status,
                "context": exam.context,
                "class_id": exam.class_id,
                "class_description": (
                    exam.class_exam.description if exam.class_exam else None
                ),
                "user_id": exam.user_id,
                "date_created": exam.date_created,
                "coordinator_id": exam.coordinator_id,
                "coordinator_full_name": (
                    f"{exam.coordinator_exam.name} {exam.coordinator_exam.surname}"
                    if exam.coordinator_exam else None
                )
            })
        return jsonify(result), 200

    except Exception as err:  # pylint: disable=broad-except
        print(err)
        return jsonify({"error": str(err)}), 500

@jwt_required()
def get_student_exams_controller():
    """
    Docstring for get_student_exams_controller
    """
    try:
        current_user_id = int(get_jwt_identity())
        exams = (
            Exam.query
            .options(
                joinedload(Exam.class_exam),         # type: ignore
                joinedload(Exam.coordinator_exam)    # type: ignore
            )
            .filter(
                Exam.date_deleted.is_(None),
                Exam.user_id == current_user_id,
                Exam.status != ExamStatus.GENERATING.value
            )
            .order_by(Exam.date_created.desc())
            .all()
        )

        result = []
        for exam in exams:
            result.append({
                "id": exam.id,
                "status": exam.status,
                "context": exam.context,
                "class_id": exam.class_id,
                "class_description": (
                    exam.class_exam.description if exam.class_exam else None
                ),
                "user_id": exam.user_id,
                "date_created": exam.date_created,
                "coordinator_id": exam.coordinator_id,
                "coordinator_full_name": (
                    f"{exam.coordinator_exam.name} {exam.coordinator_exam.surname}"
                    if exam.coordinator_exam else None
                ),
                "score": exam.score,
                "xp_gained": exam.xp_gained
            })
        return jsonify(result), 200

    except Exception as err:  # pylint: disable=broad-except
        print(err)
        return jsonify({"error": str(err)}), 500
@jwt_required()
def set_exam_on_correction(exam_id: int):
    """
    Docstring for set_exam_on_correction
    """
    try:
        current_user_id = int(get_jwt_identity())

        exam = Exam.query.get(exam_id)
        if not exam or exam.date_deleted is not None:
            return jsonify({"error": "Exam not found"}), 404

        if exam.user_id is not None and exam.user_id != current_user_id:
            return jsonify({"error": "Exam belongs to another teacher"}), 409
        if exam.status != ExamStatus.PENDING_CORRECTION.value:
            return jsonify({"error": "Exam is not Pending Correction"}), 400

        # Pasar a On Correction
        exam.status = ExamStatus.ON_CORRECTION.value
        db.session.commit()

        return jsonify({"message": "Moved to On Correction"}), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@jwt_required()
def set_exam_on_review(exam_id: int):
    """
    Docstring for set_exam_on_review
    """
    try:
        current_user_id = int(get_jwt_identity())

        exam = Exam.query.get(exam_id)
        if not exam or exam.date_deleted is not None:
            return jsonify({"error": "Exam not found"}), 404

        # Si lo tiene otra coordinadora → conflicto
        if exam.coordinator_id is not None and exam.coordinator_id != current_user_id:
            return jsonify({"error": "Exam already assigned to another coordinator"}), 409

        # Si no tiene coordinadora, asignar; si ya soy yo, mantener
        if exam.coordinator_id is None:
            exam.coordinator_id = current_user_id

        # Pasar a On Review si estaba en Pending Review
        if exam.status != ExamStatus.PENDING_REVIEW.value:
            return jsonify({"error": "Exam is not Pending Review"}), 400
        exam.status = ExamStatus.ON_REVIEW.value
        db.session.commit()

        try:
            exam_with_relations = _get_exam_with_users_and_class(exam.id)
            if exam_with_relations:
                send_exam_on_review_email_to_teacher(exam_with_relations)
        except Exception as mail_err:
            print(f"Failed to send On Review email: {mail_err}")

        return jsonify({"message": "Moved to On Review"}), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@jwt_required()
def send_exam_to_review(exam_id: int):
    """
    Docstring for send_exam_to_review
    """
    try:
        current_user_id = int(get_jwt_identity())

        exam = Exam.query.get(exam_id)
        if not exam or exam.date_deleted is not None:
            return jsonify({"error": "Exam not found"}), 404

        if exam.user_id is not None and exam.user_id != current_user_id:
            return jsonify({"error": "Exam belongs to another teacher"}), 409
        if exam.status not in (ExamStatus.GENERATING.value, ExamStatus.ON_CORRECTION.value):
            return jsonify({"error": "Exam cannot be sent to review from current status"}), 400

        # Pasar a Pending Review
        exam.status = ExamStatus.PENDING_REVIEW.value
        db.session.commit()

        return jsonify({"message": "Moved to Pending Review"}), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@jwt_required()
def accept_exam(exam_id: int):
    """
    Docstring for accept_exam
    """
    try:
        current_user_id = int(get_jwt_identity())

        exam = Exam.query.get(exam_id)
        if not exam or exam.date_deleted is not None:
            return jsonify({"error": "Exam not found"}), 404

        # Solo la coordinadora asignada puede aceptar
        if exam.coordinator_id != current_user_id:
            return jsonify({"error": "Not authorized"}), 403
        if exam.status != ExamStatus.ON_REVIEW.value:
            return jsonify({"error": "Exam is not On Review"}), 400

        exam.status = ExamStatus.ACCEPTED.value
        db.session.commit()

        try:
            exam_with_relations = _get_exam_with_users_and_class(exam.id)
            if exam_with_relations:
                send_exam_accepted_email_to_teacher(exam_with_relations)
        except Exception as mail_err:
            print(f"Failed to send accepted exam email: {mail_err}")

        return jsonify({"message": "Exam accepted"}), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


def send_to_correction(exam_id: int):
    """
    Docstring for send_to_correction
    """
    current_user_id = int(get_jwt_identity())
    exam = Exam.query.get(exam_id)
    if not exam:
        return jsonify({"message": "Exam not found"}), 404

    data = request.get_json()
    if not data:
        return jsonify({"message": "Invalid JSON body"}), 400
    if exam.coordinator_id != current_user_id:
        return jsonify({"error": "Not authorized"}), 403
    if exam.status != ExamStatus.ON_REVIEW.value:
        return jsonify({"error": "Exam is not On Review"}), 400
    exam.status = ExamStatus.PENDING_CORRECTION.value
    exam.notes = data.get("notes")

    db.session.commit()

    try:
        exam_with_relations = _get_exam_with_users_and_class(exam.id)
        if exam_with_relations:
            send_exam_sent_to_correction_email_to_teacher(exam_with_relations)
    except Exception as mail_err:
        print(f"Failed to send correction email: {mail_err}")

    return jsonify({"message": "Exam sent to correction"}), 200

@jwt_required()
def submit_correction(exam_id: int):
    """
    Persist manual teacher corrections and move the exam back to Pending Review.

    Expected payload:
    {
        "context": "updated context",
        "exercises": [
            {
                "exercise_type": "Word Formation",
                "items": [
                    {
                        "question": "Updated question",
                        "answer": "Updated answer"
                    }
                ]
            }
        ]
    }
    """
    try:
        current_user_id = int(get_jwt_identity())
        data = request.get_json(silent=True)

        if not data:
            return jsonify({"message": "Invalid JSON body"}), 400
        print("Data", data)

        edited_context = data.get("context")
        edited_exercises = data.get("exercises")

        if not isinstance(edited_context, str) or not edited_context.strip():
            return jsonify({"message": "Context is required"}), 400

        if not isinstance(edited_exercises, list) or not edited_exercises:
            return jsonify({"message": "Exercises must be a non-empty list"}), 400

        exam = (
            Exam.query
            .filter(
                Exam.id == exam_id,
                Exam.date_deleted.is_(None)
            )
            .first()
        )

        if not exam:
            return jsonify({"message": "Exam not found"}), 404
        print("Exam data", exam, exam.user_id, exam.status)

        if int(exam.user_id) != current_user_id:
            return jsonify({"message": "Unauthorized"}), 403

        if exam.status not in (
            ExamStatus.ON_CORRECTION.value,
            ExamStatus.GENERATING.value,
        ):
            return jsonify({
                "message": "Exam is not editable in current state"
            }), 400

        current_instances = cast(List[ExamExerciseInstance], exam.generated_exercises)

        print("Current", current_instances)
        if not current_instances:
            return jsonify({"message": "Exam has no generated exercises"}), 400

        if len(edited_exercises) != len(current_instances):
            return jsonify({
                "message": "You cannot change the number of exercise blocks"
            }), 400
        instances_by_type = {}
        for instance in current_instances:
            if not instance.exercise_type or not instance.exercise_type.name:
                return jsonify({"message": "Exam contains invalid exercise metadata"}), 500

            exercise_key = instance.exercise_type.name.strip().lower()

            if exercise_key in instances_by_type:
                return jsonify({
                    "message": "Duplicated exercise types are not supported"
                }), 400

            instances_by_type[exercise_key] = instance

        received_types = set()

        for exercise_block in edited_exercises:
            exercise_type_name = exercise_block.get("exercise_type")
            edited_items = exercise_block.get("items")

            if not isinstance(exercise_type_name, str) or not exercise_type_name.strip():
                return jsonify({"message": "Each exercise must include exercise_type"}), 400

            if not isinstance(edited_items, list):
                return jsonify({"message": "Each exercise must include an items list"}), 400

            exercise_key = exercise_type_name.strip().lower()

            if exercise_key in received_types:
                return jsonify({"message": "Duplicated exercise types in payload"}), 400

            instance = instances_by_type.get(exercise_key)
            if not instance:
                return jsonify({
                    "message": "You cannot change the exercise types of the exam"
                }), 400

            original_items = json.loads(instance.content_json)

            if len(edited_items) != len(original_items):
                return jsonify({
                    "message": (
                        f"You cannot change the number of items for "
                        f"'{exercise_type_name}'"
                    )
                }), 400

            updated_items = []
            updated_answers = []

            for index, (edited_item, original_item) in enumerate(
                zip(edited_items, original_items),
                start=1
            ):
                question = edited_item.get("question")
                answer = edited_item.get("answer")

                if not isinstance(question, str) or not question.strip():
                    return jsonify({
                        "message": (
                            f"Question {index} in '{exercise_type_name}' is required"
                        )
                    }), 400

                if not isinstance(answer, str) or not answer.strip():
                    return jsonify({
                        "message": (
                            f"Answer {index} in '{exercise_type_name}' is required"
                        )
                    }), 400

                updated_item = dict(original_item)
                updated_item["question"] = question.strip()
                updated_item["answer"] = answer.strip()

                updated_items.append(updated_item)
                updated_answers.append(answer.strip())

            instance.content_json = json.dumps(updated_items)
            instance.answer_key_json = json.dumps({
                "answers": updated_answers
            })

            received_types.add(exercise_key)

        exam.context = edited_context.strip()
        exam.status = ExamStatus.PENDING_REVIEW.value

        db.session.commit()

        try:
            exam_with_relations = _get_exam_with_users_and_class(exam.id)
            if exam_with_relations:
                send_exam_corrected_email_to_coordinator(exam_with_relations)
        except Exception as mail_err:
            print(f"Failed to send corrected exam email: {mail_err}")

        return jsonify({
            "message": "Exam corrected and sent back to review"
        }), 200

    except SQLAlchemyError as err:
        db.session.rollback()
        return jsonify({"message": f"Database error: {err}"}), 500

    except Exception as err:  # pylint: disable=broad-except
        db.session.rollback()
        return jsonify({"message": f"Unexpected error: {err}"}), 500
def _split_items_and_answers(items: List[dict]) -> tuple[List[dict], List[str]]:
    """
    Split generated items into:
    - public items without answer
    - answer key list

    For student exams, questions are allowed to come either:
    - with visible underscores already embedded in the question text, or
    - without underscores, in which case the frontend can render an external input.
    """

    public_items: List[dict] = []
    answers: List[str] = []

    for item in items:
        if not isinstance(item, dict):
            raise ValueError("Each item returned by the model must be an object")

        question = item.get("question")
        answer = item.get("answer")

        if not question or not isinstance(question, str):
            raise ValueError(
                "Each item returned by the model must contain a valid 'question'"
            )

        if not answer or not isinstance(answer, str):
            raise ValueError(
                "Each item returned by the model must contain a valid 'answer'"
            )

        clean_question = question.strip()
        clean_answer = answer.strip()

        if not clean_question:
            raise ValueError("Question cannot be empty")

        if not clean_answer:
            raise ValueError("Answer cannot be empty")

        answers.append(clean_answer)

        public_item = {
            "question": clean_question,
            "has_blank": "_" in clean_question,
        }

        public_items.append(public_item)

    return public_items, answers


@jwt_required()
def generate_student_exam():
    """
    Generate a student exam using:
    - class level
    - course bucket (kids/teens/adults)
    - random generic context from filesystem
    - RAG retrieval
    - LLM generation

    If the generation flow fails, a deterministic fallback exam is created.
    """

    data = request.get_json(silent=True)
    if not data:
        return jsonify({"message": "Invalid JSON body"}), 400

    class_id = data.get("class_id")
    exercise_type_ids = data.get("exercise_type_ids")
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)

    if not class_id or not exercise_type_ids:
        return jsonify({"message": "Missing required fields"}), 400

    if not isinstance(exercise_type_ids, list) or not exercise_type_ids:
        return jsonify(
            {"message": "exercise_type_ids must be a non-empty list"}
        ), 400

    if not user:
        return jsonify({"message": "Student not found"}), 404

    try:
        # ------------------------
        # Validate class
        # ------------------------
        class_obj = db.session.get(Class, class_id)
        if not class_obj or class_obj.date_deleted:
            return jsonify({"message": "Class not found or deleted"}), 404

        # ------------------------
        # Adaptive difficulty phase
        # ------------------------
        predicted_score = predict_next_student_score(
            user_id=user_id,
            class_obj=class_obj,
        )
        difficulty_band = get_difficulty_band(predicted_score)

        print(
            "Adaptive difficulty | user_id=%s | predicted_score=%s | band=%s",
            user_id,
            predicted_score,
            difficulty_band,
        )

        # ------------------------
        # Validate exercise types
        # ------------------------
        exercise_types: List[Exercise] = []

        for ex_id in exercise_type_ids:
            exercise = db.session.get(Exercise, ex_id)
            if not exercise or exercise.date_deleted:
                return jsonify(
                    {"message": f"Invalid exercise type id {ex_id}"}
                ), 400
            exercise_types.append(exercise)

        # ------------------------
        # Resolve normal context
        # ------------------------
        course_name = class_obj.description
        level = class_obj.suggested_level
        generic_context = get_random_generic_context(course_name)

        # ------------------------
        # Create exam (GENERATING)
        # ------------------------
        new_exam = Exam(
            status=ExamStatus.GENERATING.value,
            class_id=class_id,
            context=generic_context,
            user_id=user_id,
        )

        for exercise in exercise_types:
            new_exam.exercise_types.append(exercise)

        db.session.add(new_exam)
        db.session.flush()

        # ------------------------
        # Generation flow / fallback
        # ------------------------
        exercise_list_text = "\n".join(
            f"- {exercise.name}: {exercise.content_description}"
            for exercise in exercise_types
        )

        try:
            # ------------------------
            # RAG phase
            # ------------------------
            contexts = retrieve_course_context(
                course_id=course_name,
                level=level,
                exercises_description=exercise_list_text,
            )

            retrieved_context_text = "\n\n---\n\n".join(contexts)
            #print(retrieved_context_text)
            # ------------------------
            # LLM phase
            # ------------------------
            prompt = build_student_prompt(
                level=level,
                source_text=generic_context,
                exercise_list_text=exercise_list_text,
                retrieved_context=retrieved_context_text,
                difficulty_band=difficulty_band,
            )

            raw_output = generate_exam_from_llm(prompt)

            cleaned_json = extract_json(raw_output)
            parsed_output = json.loads(cleaned_json)

            exercises_output = parsed_output.get("exercises")
            if not isinstance(exercises_output, list) or not exercises_output:
                raise ValueError("Invalid exam structure returned by model")

            if len(exercises_output) != len(exercise_types):
                raise ValueError(
                    "Model did not return the expected number of exercise blocks"
                )

        except Exception as err:  # pylint: disable=broad-except
            print("Student generation flow failed, using fallback:", err)
            time.sleep(10)
            parsed_output = StudentExamFallbackService.build_exam_payload(
                level=level,
                exercise_types=exercise_types,
            )
            raw_output = StudentExamFallbackService.build_snapshot(
                reason=str(err),
                payload=parsed_output,
            )
            exercises_output = parsed_output["exercises"]

            # Override the original random bucket context with the fixed fallback one.
            new_exam.context = StudentExamFallbackService.FALLBACK_CONTEXT

        # ------------------------
        # Persist generated exercise instances
        # ------------------------
        for index, exercise_block in enumerate(exercises_output):
            instructions = exercise_block.get("instructions")
            items = exercise_block.get("items")

            if instructions is None or not isinstance(instructions, str):
                db.session.rollback()
                return jsonify({
                    "message": "Invalid exercise structure returned by model"
                }), 500

            if not isinstance(items, list) or not items:
                db.session.rollback()
                return jsonify({
                    "message": "Invalid exercise items returned by model"
                }), 500

            requested_exercise = exercise_types[index]

            final_instructions = instructions.strip()
            if not final_instructions:
                final_instructions = (
                    requested_exercise.content_description
                    or requested_exercise.name
                )

            public_items, answers = _split_items_and_answers(items)

            instance = ExamExerciseInstance(
                exam_id=new_exam.id,
                exercise_type_id=requested_exercise.id,
                instructions=final_instructions,
                content_json=json.dumps(public_items),
                answer_key_json=json.dumps({"answers": answers}),
            )
            db.session.add(instance)

        # ------------------------
        # Persist exam result
        # ------------------------
        new_exam.generated_snapshot = raw_output
        new_exam.status = ExamStatus.STUDENT_EXAM.value

        db.session.commit()

        return jsonify(
            {
                "message": "Student exam generated successfully",
                "exam_id": new_exam.id,
            }
        ), 201

    except SQLAlchemyError as err:
        db.session.rollback()
        print("db error")
        return jsonify({"message": f"Database error: {err}"}), 500

    except Exception as err:  # pylint: disable=broad-except
        db.session.rollback()
        print("unexpected error")
        return jsonify({"message": f"Unexpected error: {err}"}), 500

@jwt_required()
def submit_student_exam(exam_id: int):
    """
    Store student answers, correct the exam with the LLM,
    persist correction details, score and xp.
    """

    data = request.get_json(silent=True)
    if not data or "exercises" not in data:
        return jsonify({"message": "Invalid JSON body"}), 400

    submitted_exercises = data.get("exercises")
    if not isinstance(submitted_exercises, list) or not submitted_exercises:
        return jsonify({"message": "exercises must be a non-empty list"}), 400

    try:
        exam = (
            Exam.query
            .options(
                joinedload(Exam.class_exam),  # type: ignore
                joinedload(Exam.generated_exercises).joinedload(ExamExerciseInstance.exercise_type)  # type: ignore
            )
            .filter(
                Exam.id == exam_id,
                Exam.date_deleted.is_(None)
            )
            .first()
        )

        if not exam:
            return jsonify({"message": "Exam not found"}), 404

        if exam.status != ExamStatus.STUDENT_EXAM.value:
            return jsonify({"message": "Exam is not in resolvable status"}), 400

        instance_map = {
            instance.id: instance
            for instance in exam.generated_exercises
        }

        correction_payload = {
            "exam_id": exam.id,
            "level": exam.class_exam.suggested_level if exam.class_exam else None,
            "context": exam.context,
            "exercises": []
        }

        for submitted_block in submitted_exercises:
            instance_id = submitted_block.get("exam_exercise_instance_id")
            submitted_items = submitted_block.get("items")

            if not instance_id or not isinstance(submitted_items, list):
                db.session.rollback()
                return jsonify({"message": "Invalid submitted exercise block"}), 400

            instance = instance_map.get(instance_id)
            if not instance:
                db.session.rollback()
                return jsonify({
                    "message": f"Exercise instance {instance_id} not found"
                }), 400

            content_items = json.loads(instance.content_json or "[]")
            answer_key = json.loads(instance.answer_key_json or "{}")
            correct_answers = answer_key.get("answers", [])

            if len(submitted_items) != len(content_items) or len(submitted_items) != len(correct_answers):
                db.session.rollback()
                return jsonify({
                    "message": f"Answer count mismatch for exercise instance {instance_id}"
                }), 400

            student_answer_json = {"items": []}
            llm_items = []

            for index, submitted_item in enumerate(submitted_items):
                student_answer = (submitted_item.get("student_answer") or "").strip()
                question = content_items[index].get("question", "")
                correct_answer = correct_answers[index]

                student_answer_json["items"].append({
                    "student_answer": student_answer
                })

                llm_items.append({
                    "item_index": index,
                    "question": question,
                    "correct_answer": correct_answer,
                    "student_answer": student_answer
                })

            instance.student_answer_json = json.dumps(student_answer_json)

            correction_payload["exercises"].append({
                "exam_exercise_instance_id": instance.id,
                "exercise_type": instance.exercise_type.name if instance.exercise_type else None,
                "instructions": instance.instructions,
                "items": llm_items
            })

        prompt = build_student_correction_prompt(
            level=exam.class_exam.suggested_level if exam.class_exam else "",
            correction_payload=json.dumps(correction_payload, ensure_ascii=False)
        )

        raw_output = generate_student_correction_from_llm(prompt)

        try:
            cleaned_json = extract_json(raw_output)
            parsed_output = json.loads(cleaned_json)
        except Exception:
            db.session.rollback()
            return jsonify({"message": "Model did not return valid correction JSON"}), 500

        exercises_correction = parsed_output.get("exercises")
        model_general_feedback = (parsed_output.get("general_feedback") or "").strip()

        if not isinstance(exercises_correction, list):
            db.session.rollback()
            return jsonify({"message": "Invalid correction structure returned by model"}), 500

        model_exercises_by_id = {}
        for exercise_correction in exercises_correction:
            if not isinstance(exercise_correction, dict):
                continue

            instance_id = exercise_correction.get("exam_exercise_instance_id")
            if instance_id:
                model_exercises_by_id[instance_id] = exercise_correction

        corrected_instance_ids = set()
        normalized_exercises_correction = []
        total_items_for_score = 0
        total_weighted_points = 0.0
        full_correct_items = 0

        for expected_exercise in correction_payload["exercises"]:
            instance_id = expected_exercise["exam_exercise_instance_id"]
            instance = instance_map.get(instance_id)

            if not instance:
                db.session.rollback()
                return jsonify({"message": f"Unknown corrected instance id {instance_id}"}), 500

            model_exercise = model_exercises_by_id.get(instance_id)
            if not model_exercise:
                db.session.rollback()
                return jsonify({"message": f"Model did not correct exercise instance {instance_id}"}), 500

            model_items = model_exercise.get("items")
            if not isinstance(model_items, list):
                db.session.rollback()
                return jsonify({
                    "message": f"Invalid items structure for corrected instance {instance_id}"
                }), 500

            model_items_by_index = {}
            for model_item in model_items:
                if not isinstance(model_item, dict):
                    continue

                item_index = model_item.get("item_index")
                if isinstance(item_index, int):
                    model_items_by_index[item_index] = model_item

            normalized_items = []
            correct_count = 0
            awarded_points_sum = 0.0
            expected_items = expected_exercise.get("items", [])

            for expected_item in expected_items:
                item_index = expected_item["item_index"]
                student_answer = expected_item.get("student_answer", "")
                correct_answer = expected_item.get("correct_answer", "")
                model_item = model_items_by_index.get(item_index, {})

                raw_awarded_points = model_item.get("awarded_points")
                if raw_awarded_points is None:
                    raw_awarded_points = 1.0 if model_item.get("is_correct") is True else 0.0

                try:
                    awarded_points = float(raw_awarded_points)
                except (TypeError, ValueError):
                    awarded_points = 0.0

                if awarded_points >= 0.75:
                    awarded_points = 1.0
                elif awarded_points >= 0.25:
                    awarded_points = 0.5
                else:
                    awarded_points = 0.0

                is_correct = awarded_points == 1.0

                if is_correct:
                    correct_count += 1
                    full_correct_items += 1

                item_feedback = (model_item.get("feedback") or "").strip()
                if not item_feedback:
                    if awarded_points == 1.0:
                        item_feedback = (
                            "Correct. The answer matches the expected answer "
                            "or a clearly acceptable equivalent."
                        )
                    elif awarded_points == 0.5:
                        item_feedback = (
                            "Partially correct. The answer is related to the expected answer, "
                            "but it is incomplete or imprecise, so it received partial credit."
                        )
                    elif not student_answer:
                        item_feedback = (
                            "Incorrect. No answer was provided."
                        )
                    else:
                        item_feedback = (
                            "Incorrect. The answer does not match the expected answer closely enough."
                        )

                normalized_items.append({
                    "item_index": item_index,
                    "student_answer": student_answer,
                    "correct_answer": correct_answer,
                    "is_correct": is_correct,
                    "awarded_points": awarded_points,
                    "feedback": item_feedback
                })

                awarded_points_sum += awarded_points
                total_weighted_points += awarded_points
                total_items_for_score += 1

            total_count = len(normalized_items)
            extra_points_awarded = awarded_points_sum > correct_count

            exercise_feedback = (model_exercise.get("feedback") or "").strip()
            if not exercise_feedback:
                exercise_feedback = (
                    f"Fully correct items: {correct_count}/{total_count}. "
                    f"Weighted points: {awarded_points_sum:.1f}/{total_count}."
                )
                if extra_points_awarded:
                    exercise_feedback += (
                        " Partial credit was granted only where an answer was close "
                        "but not fully correct."
                    )

            normalized_exercise = {
                "exam_exercise_instance_id": instance_id,
                "exercise_type": expected_exercise.get("exercise_type"),
                "correct_count": correct_count,
                "total_count": total_count,
                "awarded_points_sum": awarded_points_sum,
                "feedback": exercise_feedback,
                "items": normalized_items
            }

            instance.correction_json = json.dumps(normalized_exercise)
            normalized_exercises_correction.append(normalized_exercise)
            corrected_instance_ids.add(instance_id)

        expected_instance_ids = {
            exercise["exam_exercise_instance_id"]
            for exercise in correction_payload["exercises"]
        }

        if corrected_instance_ids != expected_instance_ids:
            db.session.rollback()
            return jsonify({"message": "Model did not correct all exercise blocks"}), 500

        score = 0
        if total_items_for_score > 0:
            score = int((total_weighted_points / total_items_for_score) * 100)

        score = max(0, min(score, 100))

        general_feedback = model_general_feedback or "The exam was corrected item by item using strict grading."
        general_feedback += (
            f" Final score: {score}/100 based on "
            f"{total_weighted_points:.1f}/{total_items_for_score} weighted points "
            f"and {full_correct_items}/{total_items_for_score} fully correct answers."
        )

        if total_weighted_points > full_correct_items:
            general_feedback += (
                " Any extra points above the count of fully correct answers come only "
                "from justified partial credit reflected in the item feedback."
            )

        student = db.session.get(User, exam.user_id)
        if not student:
            db.session.rollback()
            return jsonify({"message": "Student not found for this exam"}), 404

        total_items = 0
        for instance in exam.generated_exercises:
            content = json.loads(instance.content_json or "[]")
            if isinstance(content, list):
                total_items += len(content)

        xp_gained = calculate_exam_xp(
            level=exam.class_exam.suggested_level if exam.class_exam else "",
            score=score,
            total_exercises=total_items,
        )

        new_accumulated_xp, previous_level_id, new_level_id = apply_exam_xp_to_student(
            student=student,
            xp_gained=xp_gained,
        )
        leveled_up = previous_level_id != new_level_id

        exam.student_submitted_at = datetime.datetime.now()
        exam.corrected_at = datetime.datetime.now()
        exam.score = score
        exam.xp_gained = xp_gained
        exam.llm_correction_snapshot = raw_output
        exam.score_detail_json = json.dumps({
            "general_feedback": general_feedback,
            "exercises": normalized_exercises_correction,
            "xp_gained": xp_gained,
            "student_accumulated_xp": new_accumulated_xp,
            "previous_level_id": previous_level_id,
            "new_level_id": new_level_id,
            "leveled_up": leveled_up
        })

        exam.status = ExamStatus.SOLVED.value

        db.session.commit()

        return jsonify({
            "message": "Exam submitted and corrected successfully",
            "exam_id": exam.id,
            "score": exam.score,
            "xp_gained": exam.xp_gained,
            "student_accumulated_xp": new_accumulated_xp,
            "previous_level_id": previous_level_id,
            "new_level_id": new_level_id
        }), 200

    except SQLAlchemyError as err:
        db.session.rollback()
        return jsonify({"message": f"Database error: {err}"}), 500

    except Exception as err:  # pylint: disable=broad-except
        db.session.rollback()
        return jsonify({"message": f"Unexpected error: {err}"}), 500

@jwt_required()
def get_student_exam_review(exam_id: int):
    """
    Return a corrected student exam ready for frontend review.
    """

    try:
        current_user_id = int(get_jwt_identity())

        exam = (
            Exam.query
            .options(
                joinedload(Exam.class_exam),  # type: ignore
                joinedload(Exam.generated_exercises).joinedload(ExamExerciseInstance.exercise_type),  # type: ignore
                joinedload(Exam.user_exam),  # type: ignore
            )
            .filter(
                Exam.id == exam_id,
                Exam.date_deleted.is_(None),
                Exam.user_id == current_user_id,
            )
            .first()
        )

        if not exam:
            return jsonify({"message": "Exam not found"}), 404
        if exam.status != ExamStatus.SOLVED.value:
            return jsonify({"error": "Exam is not solved yet"}), 400

        result = {
            "id": exam.id,
            "status": exam.status,
            "context": exam.context,
            "notes": exam.notes,
            "date_created": exam.date_created,
            "class_id": exam.class_id,
            "class_description": (
                exam.class_exam.description if exam.class_exam else None
            ),
            "student_id": exam.user_id,
            "student_full_name": (
                f"{exam.user_exam.name} {exam.user_exam.surname}"
                if exam.user_exam else None
            ),
            "score": exam.score,
            "xp_gained": exam.xp_gained,
            "score_detail": json.loads(exam.score_detail_json)
            if exam.score_detail_json else None,
            "exercises": []
        }

        for instance in exam.generated_exercises:
            exercise_type = instance.exercise_type

            content_items = json.loads(instance.content_json or "[]")
            student_answer_data = json.loads(instance.student_answer_json or '{"items": []}')
            correction_data = json.loads(instance.correction_json or '{"items": []}')

            student_items = student_answer_data.get("items", [])
            corrected_items = correction_data.get("items", [])

            merged_items = []

            for index, content_item in enumerate(content_items):
                merged_items.append({
                    "question": content_item.get("question"),
                    "has_blank": content_item.get("has_blank", False),
                    "student_answer": (
                        student_items[index].get("student_answer")
                        if index < len(student_items) else None
                    ),
                    "correct_answer": (
                        corrected_items[index].get("correct_answer")
                        if index < len(corrected_items) else None
                    ),
                    "is_correct": (
                        corrected_items[index].get("is_correct")
                        if index < len(corrected_items) else None
                    ),
                    "feedback": (
                        corrected_items[index].get("feedback")
                        if index < len(corrected_items) else None
                    ),
                })

            result["exercises"].append({
                "exam_exercise_instance_id": instance.id,
                "exercise_type": exercise_type.name if exercise_type else None,
                "instructions": instance.instructions,
                "feedback": correction_data.get("feedback"),
                "correct_count": correction_data.get("correct_count"),
                "total_count": correction_data.get("total_count"),
                "items": merged_items,
            })

        return jsonify(result), 200

    except Exception as err:  # pylint: disable=broad-except
        return jsonify({"message": f"Unexpected error: {err}"}), 500

def _get_exam_with_users_and_class(exam_id: int):
    return Exam.query.filter(
        Exam.id == exam_id,
        Exam.date_deleted.is_(None)
    ).first()
